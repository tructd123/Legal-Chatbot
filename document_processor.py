import os
import shutil
import PyPDF2
import pdfplumber
import fitz
from docx import Document
from pdf2image import convert_from_path
import pytesseract
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from typing import List, Tuple, Union
from tempfile import TemporaryDirectory
from PIL import ImageFilter, ImageOps

class LegalDocumentProcessor:
    def __init__(self): 
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
    
    def read_pdf(self, file_path: str, return_pages: bool = False) -> Union[str, Tuple[str, List[str]]]: 
        text_pages: list[str] = []
        missing_pages: list[int] = []
        total_pages = 0

        # --- CÁCH 1: THỬ PyMuPDF (fitz) ĐẦU TIÊN (Rất mạnh) ---
        try:
            with fitz.open(file_path) as doc:
                total_pages = len(doc)
                for idx in range(total_pages):
                    page = doc.load_page(idx)
                    page_text = page.get_text("text", sort=True) or ""
                    text_pages.append(page_text)
                    if not page_text.strip():
                        missing_pages.append(idx)
            
            if text_pages and not all(p.strip() == "" for p in text_pages):
                print(f"✅ Extracted text with PyMuPDF ({sum(len(p) for p in text_pages)} chars)")
            else:
                # PyMuPDF không lấy được gì, reset để thử cách khác
                text_pages = []
                missing_pages = []
                total_pages = 0
                raise Exception("PyMuPDF extracted no text.") # Chuyển sang Pypdfplumber
                
        except Exception as e:
            print(f"⚠️ PyMuPDF failed: {e}. Trying pdfplumber...")
            # --- HẾT CÁCH 1 ---

            # Thử pdfplumber (Code cũ của bạn, không đổi)
            try:
                with pdfplumber.open(file_path) as pdf:
                    total_pages = len(pdf.pages)
                    for idx, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        text_pages.append(page_text)
                        if not page_text.strip():
                            missing_pages.append(idx)
                if text_pages and not all(p.strip() == "" for p in text_pages):
                     print(f"✅ Extracted text with pdfplumber ({sum(len(p) for p in text_pages)} chars)")
                else:
                    raise Exception("pdfplumber extracted no text.")
            except Exception as e:
                print(f"⚠️ pdfplumber failed: {e}. Trying PyPDF2...")
                text_pages = []
                missing_pages = []
                total_pages = 0
                
                # Fallback to PyPDF2 (Code cũ của bạn, không đổi)
                try:
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        total_pages = len(reader.pages)
                        if total_pages == 0: # Xử lý file PDF bị mã hóa
                            raise Exception("PyPDF2 could not read pages (possibly encrypted).")
                        text_pages = ["" for _ in range(total_pages)] # Khởi tạo list
                        for idx, page in enumerate(reader.pages):
                            page_text = page.extract_text() or ""
                            text_pages[idx] = page_text
                            if not page_text.strip():
                                missing_pages.append(idx)
                    if text_pages and not all(p.strip() == "" for p in text_pages):
                        print(f"✅ Extracted text with PyPDF2 ({sum(len(p) for p in text_pages)} chars)")
                    else:
                         missing_pages = list(range(total_pages)) # Toàn bộ đều rỗng
                except Exception as e:
                    print(f"⚠️ PyPDF2 failed: {e}. Will attempt OCR.")
                    text_pages = []
                    missing_pages = []
                    total_pages = 0


        # --- CÁCH 2: FALLBACK TO OCR (ĐÃ TỐI ƯU HÓA) ---
        poppler_path = os.getenv("POPPLER_PATH")
        tesseract_cmd = os.getenv("TESSERACT_CMD")
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        # Nếu không có text_pages (PyMuPDF/Plumber/PDF2 đều thất bại) 
        # HOẶC có một số trang bị thiếu text
        needs_ocr = not text_pages or missing_pages

        if needs_ocr:
            print(f"ℹ️ Attempting OCR fallback...")
            try:
                # Nếu text_pages rỗng, chúng ta cần biết tổng số trang
                if total_pages == 0:
                    try:
                        # Mở lại bằng PyPDF2 chỉ để đếm trang
                        with open(file_path, 'rb') as file:
                            reader = PyPDF2.PdfReader(file)
                            total_pages = len(reader.pages)
                    except Exception:
                        # Nếu vẫn lỗi, thử đếm bằng pdfplumber
                        try:
                            with pdfplumber.open(file_path) as pdf:
                                total_pages = len(pdf.pages)
                        except Exception as e:
                            print(f"❌ Cannot determine page count for OCR: {e}")
                            return "" # Bỏ cuộc

                if not text_pages:
                    text_pages = ["" for _ in range(total_pages)]
                    missing_pages = list(range(total_pages))
                
                print(f"ℹ️ OCR-ing {len(missing_pages)} pages (one by one)...")
                
                ocr_dpi = int(os.getenv("PDF_OCR_DPI", "300")) # Giảm DPI trong .env nếu vẫn chậm
                ocr_lang = os.getenv("PDF_OCR_LANG", "vie+eng")
                ocr_config = os.getenv("PDF_OCR_CONFIG", "")

                kwargs_poppler = {"poppler_path": poppler_path} if poppler_path else {}

                verbose_ocr = os.getenv("PDF_OCR_VERBOSE", "0") == "1"

                for page_idx in missing_pages:
                    # [TỐI ƯU HÓA] Chỉ chuyển đổi 1 trang tại một thời điểm
                    try:
                        image = convert_from_path(
                            file_path,
                            dpi=ocr_dpi,
                            first_page=page_idx + 1, # pdf2image dùng index 1
                            last_page=page_idx + 1,
                            fmt="png",
                            **kwargs_poppler
                        )[0] # Lấy ảnh duy nhất trong list

                        # Tiền xử lý ảnh (code cũ của bạn)
                        img = image.convert("L")
                        img = ImageOps.equalize(img)
                        img = img.filter(ImageFilter.MedianFilter())
                        
                        page_text = pytesseract.image_to_string(
                            img,
                            lang=ocr_lang,
                            config=ocr_config
                        )
                        page_text = page_text.replace("\x0c", "").strip()
                        text_pages[page_idx] = page_text
                        if verbose_ocr:
                            char_count = len(page_text)
                            status = "chars" if char_count else "empty"
                            print(f"   ... OCR page {page_idx+1}: {char_count} {status}")
                        # print(f"   ... OCR page {page_idx+1} complete.") # Bỏ comment nếu muốn xem tiến trình
                        
                    except Exception as page_e:
                        print(f"⚠️ Failed to OCR page {page_idx+1}: {page_e}")
                        text_pages[page_idx] = "" # Đánh dấu là rỗng nếu lỗi

                print(f"✅ OCR complete. Total chars: ({sum(len(p) for p in text_pages)})")
            
            except Exception as e:
                # Khối chẩn đoán của bạn (giữ nguyên)
                pdfinfo_path = shutil.which("pdfinfo")
                print(f"⚠️ OCR fallback failed entirely: {e}")
                if not poppler_path and not pdfinfo_path:
                     print("ℹ️ Gợi ý: Cài Poppler...")
                # ... (giữ nguyên các gợi ý khác) ...

        # Kết hợp text cuối cùng
        combined_text = "\n".join(page_text for page_text in text_pages if page_text.strip())
        if combined_text.strip():
            return (combined_text, text_pages) if return_pages else combined_text

        print(f"❌ No text extracted from {file_path}. All methods failed.")
        return ("", text_pages) if return_pages else ""
    
    def read_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs).replace("\ufeff", "")
    
    def read_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read().replace("\ufeff", "")

    def process_documents(self, data_folder: str) -> List[LangchainDocument]:
        documents = []
        total_chunks = 0
        
        for filename in os.listdir(data_folder):
            file_path = os.path.join(data_folder, filename)
            ext = filename.split('.')[-1].lower()

            try: 
                if filename.endswith('.pdf'):
                    text = self.read_pdf(file_path)
                elif filename.endswith('.docx'):
                    text = self.read_docx(file_path)
                elif filename.endswith('.txt'):
                    text = self.read_txt(file_path)
                else: 
                    print(f"⚠️ Skipping unsupported file type: {filename}")
                    continue

                if not text.strip():
                    print(f"⚠️ No text extracted from {filename}, skipping.")
                    continue

                # Split into chunks
                chunks = self.text_splitter.split_text(text)
                total_chunks += len(chunks)

                for i, chunk in enumerate(chunks):
                    documents.append(LangchainDocument(
                        page_content=chunk, 
                        metadata={
                            "source": filename,
                            "chunk_index": i,
                            "document_type": ext
                        }
                    ))
                print(f"📄 Processed {filename} → {len(chunks)} chunks")
            except Exception as e:
                print(f"❌ Error processing {filename}: {e}")
        
        print(f"✅ Done! Total documents: {len(documents)} chunks across all files.")
        return documents

if __name__ == "__main__":
    processor = LegalDocumentProcessor()
    docs = processor.process_documents("data")  # Thư mục chứa PDF/DOCX/TXT
    print(f"Loaded {len(docs)} chunks.")
    if docs:
        print(docs[0].page_content[:300])  # Preview nội dung chunk đầu tiên
        print(docs[0].metadata)
    else:
        print("No documents were processed.")